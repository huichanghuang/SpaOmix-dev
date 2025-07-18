from docx import Document
from docxcompose.composer import Composer
import argparse
from pathlib import Path


def parse_args(**kwargs):
    """解析命令行参数"""
    parser = argparse.ArgumentParser(description="合并两个Word文档")
    parser.add_argument("-i","--input", help="quarto 生成的 word 文档 (report.docx)")
    parser.add_argument("-o", "--output", help="合并后的文档保存路径 (e.g., a.docx)")

    args = parser.parse_args(**kwargs)
    return args


def get_args():
    args = parse_args(
        args=["-i", "./report.docx", "-o", "./a.docx"]
    )
    return args

args = parse_args()

def delete_paragraph(paragraph):
    """删除段落"""
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


# 打开两个文档
title_page_doc = str(Path(__file__).parent / "titlepage.docx")
doc1 = Document(title_page_doc)
doc1.add_page_break()

doc2 = Document(args.input)


for para in doc2.paragraphs[:]:
    if para.style.name == "Title":
        delete_paragraph(para)
        break


composer = Composer(doc2)

composer.insert(0, doc1)


# 保存合并后的文档
composer.save(args.output)
